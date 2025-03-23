import pandas as pd
import os
from dotenv import load_dotenv, find_dotenv
import openai
from data_collection import DataCollectionAgent
from analysis import AnalysisAgent
from strategy_development import StrategyDevelopmentAgent
from portfolio_management import PortfolioManagementAgent
from report_generate import (
    generate_summary_report,
    plot_drawdown,
    plot_trade_pnl_distribution,
    plot_portfolio_value,
    plot_portfolio_with_trades
)
from datetime import datetime
from dateutil.relativedelta import relativedelta
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Initialization: Load environment variables and API key.
_ = load_dotenv(find_dotenv())
openai.api_key = os.environ['OPENAI_API_KEY']

def main():
    print("Project: ETF and Major Asset Classes Multi-Agent Trading System")
    print("Objective: Data collection, analysis, strategy & portfolio management, and report generation.\n")
    
    # Define output directory
    current_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(current_dir, "output")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Step 0: Get available ETFs and save as CSV for user reference.
    data_agent = DataCollectionAgent()
    available_etfs = data_agent.get_available_etfs()
    if not available_etfs:
        available_etfs = ["SPY", "QQQ", "IWM", "DIA", "EFA", "EEM", "VNQ"]
    etfs_df = pd.DataFrame(available_etfs, columns=['Ticker'])
    etfs_csv_file = os.path.join(output_dir, "available_etfs.csv")
    etfs_df.to_csv(etfs_csv_file, index=False)
    print("Available ETFs CSV saved at:", etfs_csv_file)
    
    # Ask the user to select ETFs.
    tickers_input = input("Please check the available ETFs CSV file and enter the tickers you want to select (comma-separated): ")
    tickers = [ticker.strip().upper() for ticker in tickers_input.split(",") if ticker.strip().upper() in available_etfs]
    if not tickers:
        print("No valid tickers selected. Defaulting to SPY, QQQ, IWM.")
        tickers = ["SPY", "QQQ", "IWM"]
    
    # Step 1: Set time window (past 5 years) and collect data.
    end_date = datetime.today()
    start_date = end_date - relativedelta(years=5)
    start_date_str = start_date.strftime("%Y-%m-%d")
    end_date_str = end_date.strftime("%Y-%m-%d")
    print(f"Time Window: {start_date_str} to {end_date_str}")
    
    portfolio_data = {}
    for ticker in tickers:
        print(f"\nCollecting data for {ticker}...")
        data = data_agent.collect_data(ticker, start_date_str, end_date_str)
        portfolio_data[ticker] = data
        ratios = data_agent.get_financial_ratios(ticker)
        if ratios:
            print(f"Financial ratios for {ticker} retrieved successfully.")
    
    risk_free_rate = data_agent.get_risk_free_rate()
    if risk_free_rate is not None:
        print("\nLatest Risk-Free Rate (10-Year Treasury):", risk_free_rate)
    
    # Step 2: Analyze data.
    analysis_agent_obj = AnalysisAgent()
    analyzed_portfolio = analysis_agent_obj.analyze_portfolio(portfolio_data)
    for ticker in tickers:
        print(f"Analysis completed for {ticker}.")
    
    # Step 3: Strategy development and portfolio management.
    strategy_agent = StrategyDevelopmentAgent()
    portfolio_agent = PortfolioManagementAgent()
    
    # Choose a sample trade date (e.g., the 11th trading day of the first selected ETF)
    sample_trade_date = analyzed_portfolio[tickers[0]].index[10]
    trade_instructions = strategy_agent.generate_trade_instructions(analyzed_portfolio, sample_trade_date)
    print("\nTrade Instructions on", sample_trade_date.date(), ":", trade_instructions)
    
    # Simulate daily portfolio updates:
    # Execute trades on the sample_trade_date; on other days, no trades are executed.
    dates = analyzed_portfolio[tickers[0]].index
    for date in dates:
        # Execute trade instructions on the sample_trade_date
        if date == sample_trade_date:
            portfolio_agent.execute_trades(trade_instructions)
        # Update portfolio value using the closing price of the first ETF
        price = analyzed_portfolio[tickers[0]].loc[date, "Close"]
        latest_prices = {tickers[0]: price}
        port_val = portfolio_agent.compute_portfolio_value(latest_prices)
        portfolio_agent.update_portfolio_history(port_val)
    
    # Convert portfolio history to a Series for plotting, using the same index as the analyzed data.
    portfolio_value_series = pd.Series(portfolio_agent.portfolio_value_history, index=dates)
    
    # Step 4: Generate report (LLM-generated analysis report)
    try:
        report_text = generate_summary_report(
            analyzed_portfolio[tickers[0]],
            additional_info="Risk-Free Rate: " + str(risk_free_rate)
        )
    except Exception as e:
        report_text = f"LLM report generation failed: {e}"
        print(report_text)
    
    # Step 5: Call plotting functions to generate four charts.
    fig_dd = plot_drawdown(portfolio_value_series)
    dd_file = os.path.join(output_dir, "drawdown.png")
    fig_dd.savefig(dd_file, dpi=100)
    plt.close(fig_dd)
    
    # Use the daily difference of portfolio values as the trade PnL distribution data.
    trade_pnl_series = portfolio_value_series.diff().dropna()
    fig_pnl_dist = plot_trade_pnl_distribution(trade_pnl_series)
    pnl_dist_file = os.path.join(output_dir, "pnl_distribution.png")
    fig_pnl_dist.savefig(pnl_dist_file, dpi=100)
    plt.close(fig_pnl_dist)
    
    fig_val = plot_portfolio_value(portfolio_value_series)
    val_file = os.path.join(output_dir, "portfolio_value.png")
    fig_val.savefig(val_file, dpi=100)
    plt.close(fig_val)
    
    # If actual trade data is available, pass it; otherwise, pass None.
    trades_data = None
    fig_trades = plot_portfolio_with_trades(portfolio_value_series, trades=trades_data)
    trades_file = os.path.join(output_dir, "portfolio_with_trades.png")
    fig_trades.savefig(trades_file, dpi=100)
    plt.close(fig_trades)
    
    # Step 6: Generate final DOCX report, including the LLM analysis report and all generated images.
    report_docx_file = os.path.join(output_dir, "report.docx")
    doc = Document()
    doc.add_heading("Strategy Report", 0)
    doc.add_paragraph(report_text)
    
    doc.add_heading("Portfolio Drawdown", level=1)
    if os.path.exists(dd_file):
        doc.add_picture(dd_file, width=Inches(6))
    
    doc.add_heading("Trade PnL Distribution", level=1)
    if os.path.exists(pnl_dist_file):
        doc.add_picture(pnl_dist_file, width=Inches(6))
    
    doc.add_heading("Portfolio Value Over Time", level=1)
    if os.path.exists(val_file):
        doc.add_picture(val_file, width=Inches(6))
    
    doc.add_heading("Portfolio Performance with Trades", level=1)
    if os.path.exists(trades_file):
        doc.add_picture(trades_file, width=Inches(6))
    
    doc.save(report_docx_file)
    print("\nGenerated Strategy Report saved at:", report_docx_file)
    
if __name__ == "__main__":
    main()